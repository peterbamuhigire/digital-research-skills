"""Witness statement builder with mandatory closers.

Per McMahon: every PI statement must end with the verbatim "free and voluntary"
+ "no threats or promises" closer to be admissible.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional


FREE_AND_VOLUNTARY_CLOSER = (
    "I have made this statement freely and voluntarily. No threats or promises "
    "have been made in connection with this statement. I have read this "
    "statement consisting of {pages} pages and find it to be true and accurate "
    "to the best of my knowledge and belief."
)


@dataclass(slots=True)
class WitnessStatement:
    case_id: str
    witness_name: str
    witness_address: str
    witness_dob: Optional[str]
    witness_phone: Optional[str]
    interviewer: str
    location: str
    started_at: datetime
    ended_at: Optional[datetime] = None
    body_paragraphs: list[str] = field(default_factory=list)
    corrections: list[str] = field(default_factory=list)
    signature_lines: list[tuple[int, str]] = field(default_factory=list)  # (page#, signed-by)


def build_statement(
    statement: WitnessStatement, *,
    output_path: str | Path,
    page_size: int = 3,  # paragraphs per "page" for closer numbering
) -> Path:
    """Render a witness statement DOCX with mandatory closers + page numbering."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Header (per McMahon's identification + qualification opener)
    h = doc.add_paragraph()
    r = h.add_run("WITNESS STATEMENT")
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(f"Case ID: {statement.case_id}")
    doc.add_paragraph(f"Witness: {statement.witness_name}")
    doc.add_paragraph(f"Address: {statement.witness_address}")
    if statement.witness_dob:
        doc.add_paragraph(f"DOB: {statement.witness_dob}")
    if statement.witness_phone:
        doc.add_paragraph(f"Phone: {statement.witness_phone}")
    doc.add_paragraph(f"Interviewer: {statement.interviewer}")
    doc.add_paragraph(f"Location: {statement.location}")
    doc.add_paragraph(f"Started: {statement.started_at.isoformat()}")
    if statement.ended_at:
        doc.add_paragraph(f"Ended: {statement.ended_at.isoformat()}")

    doc.add_paragraph()
    doc.add_paragraph("STATEMENT:")

    total_pages = max(1, (len(statement.body_paragraphs) + page_size - 1) // page_size)
    for i, para in enumerate(statement.body_paragraphs):
        doc.add_paragraph(para)
        # Insert "page N of M" + signature line at boundaries
        if (i + 1) % page_size == 0 and (i + 1) < len(statement.body_paragraphs):
            page_num = (i + 1) // page_size
            sig = doc.add_paragraph(f"Page {page_num} of {total_pages}    Signature: ____________________")
            sig.runs[0].italic = True

    if statement.corrections:
        doc.add_paragraph()
        doc.add_paragraph("Corrections (initialed by witness):")
        for c in statement.corrections:
            doc.add_paragraph(f"  • {c}")

    # Mandatory closer
    doc.add_paragraph()
    doc.add_paragraph(FREE_AND_VOLUNTARY_CLOSER.format(pages=total_pages))

    # Final signature block
    doc.add_paragraph()
    doc.add_paragraph(f"Witness signature: ____________________  Date: ____________")
    doc.add_paragraph(f"Interviewer signature: ____________________  Date: ____________")

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
