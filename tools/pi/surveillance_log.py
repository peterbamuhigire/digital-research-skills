"""Surveillance shift log with state-aware audio warnings.

McMahon: never record audio during surveillance in two-party-consent states.
This generator embeds the warning into the log header per state.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional


# Two-party (all-party) consent states (US, recording-laws status as of mid-2020s)
TWO_PARTY_CONSENT_STATES: set[str] = {
    "CA", "CT", "FL", "IL", "MD", "MA", "MI", "MT", "NV", "NH", "PA", "WA",
}


@dataclass(slots=True)
class ShiftEntry:
    timestamp: datetime
    location: str
    observation: str
    photo_refs: list[int] = field(default_factory=list)
    weather: Optional[str] = None
    lighting: Optional[str] = None


@dataclass(slots=True)
class SurveillanceLog:
    case_id: str
    subject: str
    investigator: str
    state: str           # ISO US state code, e.g. "FL"
    started_at: datetime
    location: str
    equipment: list[str] = field(default_factory=list)
    entries: list[ShiftEntry] = field(default_factory=list)

    @property
    def audio_recording_lawful(self) -> bool:
        """Returns False if state is two-party-consent — embed in log header."""
        return self.state.upper() not in TWO_PARTY_CONSENT_STATES

    def to_docx(self, output_path: str | Path) -> Path:
        from docx import Document
        from docx.shared import Pt, RGBColor

        doc = Document()
        h = doc.add_paragraph()
        r = h.add_run("SURVEILLANCE SHIFT LOG")
        r.bold = True
        r.font.size = Pt(14)

        doc.add_paragraph(f"Case ID: {self.case_id}")
        doc.add_paragraph(f"Subject: {self.subject}")
        doc.add_paragraph(f"Investigator: {self.investigator}")
        doc.add_paragraph(f"Location: {self.location}")
        doc.add_paragraph(f"State: {self.state}")
        doc.add_paragraph(f"Started: {self.started_at.isoformat()}")
        doc.add_paragraph(f"Equipment: {', '.join(self.equipment)}")

        # Audio-recording warning
        warn = doc.add_paragraph()
        if not self.audio_recording_lawful:
            r = warn.add_run(
                f"⚠ {self.state} is a TWO-PARTY-CONSENT state. "
                "DO NOT record audio without all-party consent — criminal penalty."
            )
            r.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        else:
            r = warn.add_run(
                f"{self.state} is a one-party-consent state. "
                "Audio recording lawful when investigator is a party to the conversation."
            )
            r.italic = True

        # Entries table
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "Time"
        hdr[1].text = "Location"
        hdr[2].text = "Observation"
        hdr[3].text = "Weather/Light"
        hdr[4].text = "Photo refs"
        for e in self.entries:
            row = table.add_row().cells
            row[0].text = e.timestamp.isoformat()
            row[1].text = e.location
            row[2].text = e.observation
            row[3].text = f"{e.weather or '-'} / {e.lighting or '-'}"
            row[4].text = ", ".join(str(p) for p in e.photo_refs)

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path
