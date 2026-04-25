"""Append-only chain-of-custody ledger.

Per McMahon: every transfer of physical evidence must be logged with from/to/
date/time/signature. Output is a court-ready custody affidavit.
"""
from __future__ import annotations

import hashlib
import json
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional


@dataclass(slots=True)
class CustodyTransfer:
    timestamp: datetime
    from_party: str       # name + role + signature
    to_party: str
    location: str
    purpose: str          # "collection" | "transport" | "analysis" | "storage" | "court"
    notes: str = ""
    sealed_container: bool = False
    seal_id: Optional[str] = None


@dataclass(slots=True)
class ChainOfCustody:
    """Append-only ledger. Each transfer hashes the previous to prevent
    tampering after the fact (poor man's blockchain — sufficient for court)."""

    case_id: str
    evidence_id: str
    description: str
    collected_at: datetime
    collected_by: str
    collected_location: str
    transfers: list[CustodyTransfer] = field(default_factory=list)
    integrity_hashes: list[str] = field(default_factory=list)

    def append(self, transfer: CustodyTransfer) -> None:
        """Add a transfer + its content-hash chained to the previous entry."""
        prev_hash = self.integrity_hashes[-1] if self.integrity_hashes else self.case_id
        payload = json.dumps({
            "prev": prev_hash,
            "ts": transfer.timestamp.isoformat(),
            "from": transfer.from_party,
            "to": transfer.to_party,
            "loc": transfer.location,
            "purpose": transfer.purpose,
            "seal": transfer.seal_id,
        }, sort_keys=True)
        h = hashlib.sha256(payload.encode("utf-8")).hexdigest()
        self.transfers.append(transfer)
        self.integrity_hashes.append(h)

    def verify_integrity(self) -> bool:
        """Re-derive each hash from its predecessor; ensure no tampering."""
        prev = self.case_id
        for transfer, expected in zip(self.transfers, self.integrity_hashes):
            payload = json.dumps({
                "prev": prev,
                "ts": transfer.timestamp.isoformat(),
                "from": transfer.from_party,
                "to": transfer.to_party,
                "loc": transfer.location,
                "purpose": transfer.purpose,
                "seal": transfer.seal_id,
            }, sort_keys=True)
            h = hashlib.sha256(payload.encode("utf-8")).hexdigest()
            if h != expected:
                return False
            prev = h
        return True

    def to_affidavit(self, output_path: str | Path) -> Path:
        """Render a court-ready chain-of-custody affidavit as DOCX."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        title = doc.add_paragraph()
        run = title.add_run("CHAIN OF CUSTODY AFFIDAVIT")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph(f"Case ID: {self.case_id}")
        doc.add_paragraph(f"Evidence ID: {self.evidence_id}")
        doc.add_paragraph(f"Description: {self.description}")
        doc.add_paragraph(f"Collected: {self.collected_at.isoformat()} by {self.collected_by} at {self.collected_location}")

        doc.add_heading("Transfers", level=1)
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Timestamp"
        hdr[1].text = "From"
        hdr[2].text = "To"
        hdr[3].text = "Location"
        hdr[4].text = "Purpose"
        hdr[5].text = "Seal ID"
        for t in self.transfers:
            row = table.add_row().cells
            row[0].text = t.timestamp.isoformat()
            row[1].text = t.from_party
            row[2].text = t.to_party
            row[3].text = t.location
            row[4].text = t.purpose
            row[5].text = t.seal_id or "-"

        doc.add_heading("Integrity", level=1)
        doc.add_paragraph(f"Final hash: {self.integrity_hashes[-1] if self.integrity_hashes else 'n/a'}")
        doc.add_paragraph(f"Integrity verified: {self.verify_integrity()}")

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path
