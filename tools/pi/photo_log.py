"""Photo / video log generator — McMahon's Figures 13.1-13.3 in code.

Auto-fills from EXIF where available; produces a court-ready exhibit log.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional


@dataclass(slots=True)
class PhotoEntry:
    photo_number: int
    filename: str
    captured_at: Optional[datetime] = None
    photographer: str = ""
    camera: str = ""
    lens: str = ""
    iso: Optional[int] = None
    aperture: Optional[str] = None
    shutter: Optional[str] = None
    flash: bool = False
    location: str = ""
    description: str = ""
    exhibit_label: Optional[str] = None  # "Exhibit A-3" etc.


@dataclass(slots=True)
class PhotoLog:
    case_id: str
    investigator: str
    location: str
    date: datetime
    entries: list[PhotoEntry] = field(default_factory=list)

    def add_from_exif(self, image_path: str | Path, description: str = "") -> PhotoEntry:
        """Append an entry auto-populated from the image's EXIF."""
        from ..verification.exif import extract_exif

        data = extract_exif(str(image_path))
        captured_at = None
        if data.captured_at:
            try:
                captured_at = datetime.fromisoformat(str(data.captured_at).replace(":", "-", 2))
            except ValueError:
                captured_at = None
        entry = PhotoEntry(
            photo_number=len(self.entries) + 1,
            filename=Path(image_path).name,
            captured_at=captured_at,
            photographer=self.investigator,
            camera=data.camera_model or "",
            description=description,
        )
        self.entries.append(entry)
        return entry

    def to_docx(self, output_path: str | Path) -> Path:
        """Render the log as a DOCX exhibit."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        h = doc.add_paragraph()
        r = h.add_run("PHOTO / VIDEO LOG")
        r.bold = True
        r.font.size = Pt(14)

        doc.add_paragraph(f"Case ID: {self.case_id}")
        doc.add_paragraph(f"Investigator: {self.investigator}")
        doc.add_paragraph(f"Location: {self.location}")
        doc.add_paragraph(f"Date: {self.date.isoformat()}")

        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "#"
        hdr[1].text = "Filename"
        hdr[2].text = "Captured"
        hdr[3].text = "Camera"
        hdr[4].text = "Description"
        hdr[5].text = "Exhibit"
        for e in self.entries:
            row = table.add_row().cells
            row[0].text = str(e.photo_number)
            row[1].text = e.filename
            row[2].text = e.captured_at.isoformat() if e.captured_at else "-"
            row[3].text = e.camera
            row[4].text = e.description
            row[5].text = e.exhibit_label or "-"

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path
