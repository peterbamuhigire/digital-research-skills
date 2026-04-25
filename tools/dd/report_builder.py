"""Hetherington-style DD report builder.

Renders the canonical DD report skeleton: cover (Privileged & Confidential
+ Attorney Work Product where applicable), executive summary FIRST with
red/amber/green flags, body in priority order, fixed liability disclaimer.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Optional


# Hetherington's verbatim liability-management language. Adapt only with legal review.
LIABILITY_DISCLAIMER = (
    "DISCLAIMER: The information contained in this report is based on public records, "
    "commercial databases, and other open sources available at the time of preparation. "
    "While reasonable steps have been taken to verify the accuracy of the data, errors "
    "and omissions can occur in source databases, and information may have changed since "
    "compilation. This report is provided for the exclusive use of the named recipient "
    "and may not be relied upon by any other party. The investigator makes no warranty, "
    "express or implied, regarding the accuracy, completeness, or fitness for any "
    "particular purpose of the information contained herein."
)


@dataclass(slots=True)
class FlaggedFinding:
    flag: str            # "red" | "amber" | "green"
    title: str
    summary: str
    sources: list[str] = field(default_factory=list)


@dataclass(slots=True)
class DDReport:
    client: str
    subject_name: str
    subject_type: str    # "person" | "company"
    objective: str       # one-paragraph purpose
    privileged: bool = True
    attorney_work_product: bool = False
    investigator_firm: str = "digital-research-engine"
    report_date: date = field(default_factory=date.today)

    executive_summary: str = ""
    flags: list[FlaggedFinding] = field(default_factory=list)
    cara_indicators: dict[str, str] = field(default_factory=dict)  # for persons
    swot: dict[str, list[str]] = field(default_factory=dict)        # for companies

    body_sections: list[tuple[str, str]] = field(default_factory=list)  # (title, content)
    sources_cited: list[str] = field(default_factory=list)


def build_dd_report(report: DDReport, *, output_path: str | Path) -> Path:
    """Render the report as a Word document at `output_path`."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DUE DILIGENCE REPORT\n")
    run.bold = True
    run.font.size = Pt(20)

    if report.privileged:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("PRIVILEGED AND CONFIDENTIAL")
        r2.bold = True
        r2.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    if report.attorney_work_product:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("ATTORNEY WORK PRODUCT")
        r3.bold = True
        r3.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()
    doc.add_paragraph(f"Subject: {report.subject_name}")
    doc.add_paragraph(f"Client: {report.client}")
    doc.add_paragraph(f"Prepared by: {report.investigator_firm}")
    doc.add_paragraph(f"Report date: {report.report_date.isoformat()}")

    doc.add_page_break()

    # Objective
    doc.add_heading("Objective", level=1)
    doc.add_paragraph(report.objective)

    # Executive summary with flags FIRST (Hetherington rule)
    doc.add_heading("Executive Summary", level=1)
    if report.executive_summary:
        doc.add_paragraph(report.executive_summary)

    if report.flags:
        doc.add_heading("Key Findings — Traffic Light", level=2)
        for f in sorted(report.flags, key=lambda x: {"red": 0, "amber": 1, "green": 2}[x.flag]):
            p = doc.add_paragraph()
            r = p.add_run(f"[{f.flag.upper()}] {f.title}")
            r.bold = True
            colours = {"red": RGBColor(0xC0, 0x00, 0x00),
                       "amber": RGBColor(0xC0, 0x80, 0x00),
                       "green": RGBColor(0x00, 0x80, 0x00)}
            r.font.color.rgb = colours[f.flag]
            doc.add_paragraph(f.summary)
            if f.sources:
                doc.add_paragraph(f"Sources: {', '.join(f.sources)}").italic = True

    # CARA (for persons)
    if report.subject_type == "person" and report.cara_indicators:
        doc.add_heading("CARA Profile", level=1)
        for k, v in report.cara_indicators.items():
            doc.add_paragraph(f"{k}: {v}")

    # SWOT (for companies)
    if report.subject_type == "company" and report.swot:
        doc.add_heading("SWOT Summary", level=1)
        for k in ("strengths", "weaknesses", "opportunities", "threats"):
            if k in report.swot:
                doc.add_heading(k.title(), level=2)
                for item in report.swot[k]:
                    doc.add_paragraph(item, style="List Bullet")

    # Body sections (priority order — caller orders them)
    for title, content in report.body_sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)

    # Sources cited
    if report.sources_cited:
        doc.add_heading("Sources Cited", level=1)
        for s in report.sources_cited:
            doc.add_paragraph(s, style="List Bullet")

    # Disclaimer (always last)
    doc.add_page_break()
    doc.add_heading("Disclaimer", level=1)
    p = doc.add_paragraph(LIABILITY_DISCLAIMER)
    p.runs[0].font.size = Pt(9)
    p.runs[0].italic = True

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
