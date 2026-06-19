from __future__ import annotations

import re
import shutil
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass
from pathlib import Path

import xlsxwriter
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PROJECT = ROOT / "projects" / "healthcare-app-clinical-data"
OUTPUT = PROJECT / "05-output" / "medic8-global-settings"
EXPORT = PROJECT / "export" / "medic8-global-settings"
DATE = "2026-05-06"
VERSION = "v2"
SEP_RE = re.compile(r"^\|\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")


@dataclass(frozen=True)
class Cohort:
    slug: str
    title: str
    purpose: str
    primary_key: str
    category_field: str
    standards: tuple[str, ...]
    import_target: str
    preconfigure_scope: str
    facility_confirmation: str
    owner: str


COHORTS: tuple[Cohort, ...] = (
    Cohort("tenant-blueprints", "Tenant Blueprints", "Canonical setup recipes that choose modules, roles, workflows, forms, reports, inventory, prices, and integrations by facility type.", "blueprint_id", "facility_type", ("Country packs", "Facility master", "RBAC", "Workflow engine"), "tenant_blueprints / provisioning runner", "Preconfigure as selectable setup scripts.", "Facility confirms country, facility type, services, and operating model.", "Medic8 implementation lead"),
    Cohort("country-packs", "Country Packs", "Localisation constants for countries, currencies, timezones, regulators, privacy law, forms, and default blueprints.", "country_code", "country_name", ("ISO 3166", "ISO 4217", "IANA timezone", "National privacy laws", "National regulators"), "global_country_profiles", "Preconfigure all full country packs; fail closed for stubs.", "Facility confirms country and reporting identifiers.", "Platform admin"),
    Cohort("facilities", "Facilities / Organisation Master", "Facility type and tier defaults used to select modules, capabilities, referral rules, and blueprint options.", "facility_type_id", "country", ("National facility registries", "Health-sector facility tier systems", "GS1 GLN"), "global_facility_types / facility_profiles", "Preconfigure facility-type templates.", "Facility confirms legal identity, licence, address, national facility code, and services offered.", "Facility director"),
    Cohort("roles-permissions", "Roles, Cadres, and Permissions", "Cadre-aligned roles, regulator mapping, and permission bundles for safe default user creation.", "role_id", "role_family", ("Professional councils", "RBAC", "Least privilege"), "global_roles / role_permission_map", "Preconfigure role templates and permission bundles.", "Facility supplies named staff and licence evidence.", "Facility admin / HR"),
    Cohort("workflows", "Workflows and Care Pathways", "Default operational and clinical workflows that become order sets, queues, forms, reports, and smoke tests.", "workflow_id", "workflow_domain", ("National clinical guidelines", "HMIS/DHIS2", "IDSR"), "workflow_templates / order_sets", "Preconfigure by blueprint and country.", "Facility confirms which workflows are active and who owns them.", "Medical director"),
    Cohort("conditions", "Conditions", "Diagnosis/problem-list seed catalogue with reporting and clinical terminology mapping.", "icd10_code", "icd10_chapter", ("ICD-10", "ICD-11", "SNOMED CT", "Uganda IDSR"), "diagnosis_catalogue / problem_list_defaults", "Preconfigure global diagnosis catalogue and top problem-list subsets.", "Medical director confirms active problem list and specialty add-ons.", "Medical director"),
    Cohort("drugs", "Drugs", "Drug formulary seed catalogue, ATC/DDD/RxNorm mapping, LASA flags, and local formulary references.", "atc_code", "atc_level1", ("ATC/DDD", "WHO EML", "Uganda EMHSLU", "KEML", "RxNorm", "NDA/PPB/TMDA"), "drug_catalogue / formulary_defaults", "Preconfigure global medicine catalogue and tiered starter formularies.", "Pharmacist confirms stocked formulary, prices, suppliers, and licence restrictions.", "Pharmacist-in-charge"),
    Cohort("drug-interactions", "Drug-Drug Interactions", "Clinical decision support interaction pairs and severity handling.", "interaction_id", "severity", ("DDInter", "ATC", "CDS safety"), "drug_interaction_rules", "Preconfigure CDS rule base.", "Clinical safety lead confirms alert display and override policy.", "Clinical safety lead"),
    Cohort("paediatric-dosing", "Paediatric Dosing", "Paediatric and neonatal dosing rules linked to drugs, units, age/weight bands, and safety checks.", "rule_id", "age_band", ("WHO EMLc", "WHO Children's Formulary", "ATC", "UCUM"), "paediatric_dosing_rules", "Preconfigure dosing rules where source-backed.", "Clinician/pharmacist validates local protocol and high-risk neonatal rules.", "Paediatric clinical lead"),
    Cohort("allergens", "Allergens", "Drug, food, and environmental allergen master for patient safety and CDS.", "allergen_id", "allergen_type", ("SNOMED CT", "RxNorm", "AllergyIntolerance"), "allergen_master", "Preconfigure global allergen master.", "Facility confirms allergy capture workflow and any local allergen labels.", "Clinical safety lead"),
    Cohort("lab-tests", "Laboratory Tests", "LOINC-coded lab menu with specimen, container, units, reference ranges, TAT, critical values, and BOM linkage.", "loinc_code", "loinc_class", ("LOINC", "UCUM", "WHO EDL", "ISO 15189"), "lab_test_catalogue", "Preconfigure global lab catalogue and blueprint-specific menus.", "Lab manager confirms in-house menu, analyser method, reference ranges, prices, and accreditation status.", "Lab manager"),
    Cohort("ucum", "UCUM Units", "Canonical units and conversion edges used across labs, pharmacy, dosing, inventory, and billing.", "from_unit", "dimension", ("UCUM", "ISO 80000"), "unit_master / conversion_graph", "Preconfigure globally before dependent catalogues import.", "Facility only confirms display preferences where needed.", "Platform admin"),
    Cohort("imaging", "Imaging", "Imaging/radiology catalogue with modality, body region, DICOM/RadLex mapping, reports, and BOM linkage.", "loinc_code", "modality", ("DICOM", "RadLex", "LOINC", "ACR"), "imaging_catalogue", "Preconfigure global imaging catalogue and modality subsets.", "Radiology lead confirms installed modalities, PACS details, prices, and contrast stock.", "Radiographer / imaging in-charge"),
    Cohort("procedures", "Procedures and Services", "Procedure, intervention, service, consent, role, and billing defaults.", "procedure_name", "category", ("ICHI", "ICD-10-PCS", "CDT", "WHO Surgical Safety Checklist"), "service_catalogue / procedure_catalogue", "Preconfigure global procedure/service catalogue.", "Medical director and accountant confirm active services, prices, revenue accounts, and consent requirements.", "Medical director + accountant"),
    Cohort("consumables", "Consumables and Supplies", "Medical supplies, devices, reagents, pack sizes, storage, and procurement classification.", "item_id", "primary_category", ("GMDN", "UNSPSC", "GS1", "Uganda EMHSLU", "ISO 13485"), "item_master / inventory_item_defaults", "Preconfigure global item master and starter inventory lists.", "Store keeper confirms physical stock, suppliers, unit costs, lots, and expiry dates.", "Store keeper"),
    Cohort("boms", "Bills of Materials", "Default kits and auto-deduction line items for services, lab tests, imaging, vaccines, and procedures.", "bom_id", "bom_family", ("BOM", "WHO PQS", "LOINC", "Procedure links"), "clinical_boms / auto_deduction_rules", "Preconfigure default BOMs and require approval for active events.", "Domain owner confirms substitutions, pack quantities, and facility SOP differences.", "Pharmacist / lab manager / theatre lead"),
    Cohort("vaccines", "Vaccines", "Vaccine catalogue, schedules, cold-chain, MDVP, and AEFI defaults.", "vaccine_id", "antigen", ("ATC J07", "WHO EPI", "WHO PQS", "Brighton AEFI"), "vaccine_catalogue / immunisation_schedule", "Preconfigure country schedules and EPI-enabled blueprint defaults.", "EPI focal person confirms stocked antigens, cold-chain location, and schedule variants.", "EPI focal person"),
    Cohort("standard-forms", "Standard Forms", "Paper forms, registers, cards, reports, and DHIS2/HMIS equivalents.", "form_code", "domain", ("Uganda HMIS", "IDSR", "DHIS2", "mTrac/eIDSR", "OpenMRS"), "form_registry / report_templates", "Preconfigure mandatory country/blueprint form sets.", "Facility confirms reporting unit IDs and optional/private forms.", "Records officer / M&E lead"),
    Cohort("reporting-kpis", "Reporting KPIs", "Indicator library, numerator/denominator definitions, dashboards, and report mappings.", "indicator_id", "indicator_domain", ("WHO 100 Core", "HMIS", "PEPFAR MER", "Global Fund", "DHIS2"), "kpi_registry / dashboard_defaults", "Preconfigure KPI definitions and blueprint dashboards.", "M&E lead confirms mandatory reports, donor IDs, thresholds, and reporting calendar.", "M&E lead"),
    Cohort("billing-tariffs", "Billing, Tariffs, and Insurance", "Charge items, price-list templates, payer mappings, public tariffs, SHA/insurance defaults.", "charge_item_id", "charge_category", ("ISO 4217", "SHA", "IRA", "IFRS revenue mapping"), "billing_catalogue / price_lists / payer_maps", "Preconfigure charge items and source-backed tariff skeletons.", "Accountant confirms prices, insurer contracts, tax treatment, and revenue accounts.", "Accountant"),
    Cohort("holiday-calendars", "Holiday Calendars", "Country holiday defaults for appointments, payroll, reporting deadlines, and operational calendars.", "holiday_id", "country_code", ("Country public holiday law", "IANA timezone"), "holiday_calendar", "Preconfigure five-year country calendars.", "Facility confirms local closures and exceptional working days.", "Facility admin"),
)


STANDARDS = {
    "ICD-10": ("World Health Organization", "Diagnosis reporting and morbidity classification.", "WHO ICD-10 2019", "Global diagnosis coding baseline."),
    "ICD-11": ("World Health Organization", "Forward path for diagnosis coding.", "Current WHO ICD-11", "Future candidate mapping."),
    "SNOMED CT": ("SNOMED International", "Clinical concept terminology; licence-sensitive.", "Current International Edition", "Clinical semantics and FHIR bridge."),
    "Uganda IDSR": ("Uganda Ministry of Health / WHO AFRO", "Notifiable disease surveillance.", "Uganda IDSR guidelines", "Reporting and alert workflows."),
    "ATC/DDD": ("WHO Collaborating Centre for Drug Statistics Methodology", "Drug classification and utilisation unit.", "Current WHO ATC/DDD", "Medication hierarchy and DDD fields."),
    "WHO EML": ("World Health Organization", "Essential medicines recommendation.", "2025 WHO EML/EMLc page archived", "International essential-medicine comparator."),
    "Uganda EMHSLU": ("Uganda Ministry of Health", "Uganda public-sector medicines and supplies reference.", "EMHSLU 2023 in corpus", "Local default formulary and supply list."),
    "KEML": ("Kenya Ministry of Health", "Kenya essential medicines reference.", "KEML 2023 in corpus", "Kenya formulary triangulation."),
    "RxNorm": ("US National Library of Medicine", "Medication normalisation and API reference.", "NLM RxNorm overview archived", "Interoperability bridge."),
    "NDA/PPB/TMDA": ("National medicines regulators", "Medicine registration and pharmacy/drug-shop regulation.", "Official regulator sources in country packs", "Legal marketing and licence filters."),
    "DDInter": ("DDInter dataset maintainers", "Drug interaction dataset.", "DDInter v2 imported in corpus", "CDS interaction pairs."),
    "CDS safety": ("Medic8 clinical safety governance", "Runtime clinical decision support governance.", "Project implementation rule", "Alert severity and override design."),
    "WHO EMLc": ("World Health Organization", "Children's essential medicines.", "WHO EMLc / children's formulary sources", "Paediatric medicine baseline."),
    "WHO Children's Formulary": ("World Health Organization", "Paediatric dose guidance.", "Sources in cohort", "Age/weight dosing."),
    "UCUM": ("UCUM Organization / Regenstrief", "Canonical units.", "Official UCUM site archived", "Unit normalisation and conversion."),
    "ISO 80000": ("International Organization for Standardization", "Quantities and units reference.", "Current ISO family", "Dimensional sanity checks."),
    "LOINC": ("Regenstrief Institute", "Lab and observation codes.", "LOINC licence archived", "Lab and observation identifiers."),
    "WHO EDL": ("World Health Organization", "Essential diagnostics reference.", "WHO EDL cited in corpus", "Diagnostic menu comparator."),
    "ISO 15189": ("International Organization for Standardization", "Medical laboratory quality and competence.", "Current ISO standard", "Lab accreditation flag."),
    "DICOM": ("DICOM Standards Committee", "Medical imaging exchange.", "DICOM site archived", "PACS/imaging integration."),
    "RadLex": ("Radiological Society of North America", "Radiology terminology.", "RadLex sources in corpus", "Imaging anatomy/procedure terms."),
    "ACR": ("American College of Radiology", "Imaging guidance.", "ACR sources in corpus", "Imaging appropriateness and protocols."),
    "ICHI": ("World Health Organization", "Health intervention classification.", "WHO ICHI references in corpus", "Open procedure coding baseline."),
    "ICD-10-PCS": ("CMS/NCHS", "US inpatient procedure coding.", "Edition cited in corpus", "Secondary procedure crosswalk."),
    "CDT": ("American Dental Association", "Dental code set; licence-sensitive.", "CDT rows in corpus", "Dental procedure mapping."),
    "WHO Surgical Safety Checklist": ("World Health Organization", "Surgical safety workflow.", "WHO checklist", "Theatre/process defaults."),
    "GMDN": ("GMDN Agency", "Medical device nomenclature.", "GMDN references in corpus", "Device identity."),
    "UNSPSC": ("UNSPSC", "Procurement taxonomy.", "UNSPSC references in corpus", "Supply-chain grouping."),
    "GS1": ("GS1", "Trade item/location identifiers.", "GS1 GTIN/GLN", "Supply and facility identifiers."),
    "ISO 13485": ("International Organization for Standardization", "Medical-device quality systems.", "Current ISO standard", "Quality context for device vendors."),
    "WHO PQS": ("World Health Organization", "Vaccine cold-chain/injection equipment specifications.", "PQS references in corpus", "Cold-chain and AD-syringe defaults."),
    "ATC J07": ("WHO ATC/DDD", "Vaccine ATC class.", "ATC J07", "Vaccine classification."),
    "WHO EPI": ("World Health Organization / national EPI", "Immunisation schedule design.", "Country schedule sources", "Schedule defaults."),
    "Brighton AEFI": ("Brighton Collaboration", "AEFI case definitions.", "Sources in cohort", "Adverse-event categorisation."),
    "Uganda HMIS": ("Uganda Ministry of Health", "Routine reporting forms.", "Cached HMIS sources", "Form/report defaults."),
    "IDSR": ("Uganda MoH / WHO AFRO", "Disease surveillance.", "Cached IDSR source", "Alert/report forms."),
    "DHIS2": ("HISP / national MoH deployments", "Aggregate reporting platform.", "Country pack sources", "Reporting endpoint target."),
    "mTrac/eIDSR": ("Uganda Ministry of Health", "Electronic surveillance path.", "Uganda IDSR/HMIS sources", "Weekly surveillance."),
    "OpenMRS": ("OpenMRS community", "HIV/ART longitudinal EMR reference.", "Uganda HIV guideline cache", "Programme interoperability."),
    "WHO 100 Core": ("World Health Organization", "Core health indicators.", "Cohort sources", "KPI comparator."),
    "HMIS": ("National ministries of health", "Routine health reporting.", "Country pack/form sources", "Numerator/denominator mapping."),
    "PEPFAR MER": ("PEPFAR", "HIV programme reporting.", "MER references in corpus", "HIV/DATIM indicators."),
    "Global Fund": ("Global Fund", "Programme monitoring framework.", "Cohort sources", "Grant reporting."),
    "ISO 4217": ("International Organization for Standardization", "Currency codes.", "ISO page archived", "Billing/currency defaults."),
    "SHA": ("Kenya Social Health Authority", "Kenya health financing claims.", "Kenya country/billing sources", "Payer mapping."),
    "IRA": ("Insurance regulators", "Insurance oversight.", "Country pack sources", "Insurer registry context."),
    "IFRS revenue mapping": ("IFRS Foundation / accounting policy", "Revenue treatment.", "Accounting design docs", "GL mapping."),
    "Country public holiday law": ("National governments", "Public holiday dates.", "Country sources in cohort", "Calendar defaults."),
    "IANA timezone": ("IANA", "Timezone identifiers.", "Country pack sources", "Scheduling/local time."),
    "ISO 3166": ("International Organization for Standardization", "Country/subdivision codes.", "Country pack sources", "Country identity."),
    "National privacy laws": ("National legislatures/data protection authorities", "Personal-data processing rules.", "Country pack sources", "Privacy defaults."),
    "National regulators": ("Health-sector statutory bodies", "Professional/facility/medicine regulation.", "Country pack and roles sources", "Licence evidence fields."),
    "National facility registries": ("Ministries of health", "Facility registry identifiers.", "Country pack/facility sources", "Facility code validation."),
    "Health-sector facility tier systems": ("Ministries of health", "Facility level and service scope.", "Facility cohort sources", "Blueprint selection."),
    "GS1 GLN": ("GS1", "Location identifier.", "Onboarding contract", "Optional facility/location identifier."),
    "RBAC": ("Medic8 platform policy", "Role-based access control.", "Local app sources", "Permission seeding."),
    "Least privilege": ("Security control principle", "Deny-by-default access.", "Project implementation standard", "Permission defaults."),
    "Country packs": ("Project integration corpus", "Country constants.", "Project cohort", "Blueprint dependency."),
    "Facility master": ("Project integration corpus", "Facility type definitions.", "Project cohort", "Blueprint dependency."),
    "Workflow engine": ("Medic8 platform", "Provisionable workflow templates.", "Local implementation target", "Onboarding automation."),
    "Professional councils": ("Country statutory councils", "Cadre registration.", "Roles/country sources", "Licence evidence."),
    "National clinical guidelines": ("Ministries of health", "Care pathways.", "Workflow sources", "Order-set defaults."),
    "AllergyIntolerance": ("HL7 FHIR", "Allergy resource model.", "FHIR standard", "Interop target."),
    "BOM": ("Medic8 inventory model", "Bill of materials / auto-deduction.", "Project cohort", "Stock decrement logic."),
    "Procedure links": ("Cross-cohort references", "Procedure to BOM linkage.", "Project corpus", "Auto-deduction relation."),
}


def read(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")


def clean_cell(value: str) -> str:
    value = re.sub(r"<br\s*/?>", "\n", value, flags=re.I)
    value = re.sub(r"`([^`]*)`", r"\1", value)
    value = value.replace("**", "")
    return (
        value.strip()
        .replace("&#124;", "|")
        .replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
    )


def split_row(line: str) -> list[str]:
    row = line.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    cells: list[str] = []
    cur: list[str] = []
    escaped = False
    for ch in row:
        if escaped:
            cur.append(ch)
            escaped = False
            continue
        if ch == "\\":
            escaped = True
            continue
        if ch == "|":
            cells.append(clean_cell("".join(cur)))
            cur = []
        else:
            cur.append(ch)
    cells.append(clean_cell("".join(cur)))
    return cells


def markdown_tables(path: Path):
    lines = read(path).splitlines()
    i = 0
    while i < len(lines) - 1:
        if lines[i].strip().startswith("|") and SEP_RE.match(lines[i + 1].strip()):
            headers = split_row(lines[i])
            rows = []
            j = i + 2
            while j < len(lines):
                stripped = lines[j].strip()
                if not stripped:
                    k = j + 1
                    while k < len(lines) and not lines[k].strip():
                        k += 1
                    if (
                        k < len(lines)
                        and lines[k].strip().startswith("|")
                        and not (k + 1 < len(lines) and SEP_RE.match(lines[k + 1].strip()))
                    ):
                        j = k
                        continue
                    break
                if not stripped.startswith("|"):
                    break
                if j + 1 < len(lines) and SEP_RE.match(lines[j + 1].strip()):
                    break
                rows.append((lines[j], split_row(lines[j])))
                j += 1
            yield headers, rows
            i = j
        else:
            i += 1


def normalize(headers: list[str], values: list[str], raw_line: str) -> dict[str, str] | None:
    if not values or SEP_RE.match(raw_line.strip()):
        return None
    if values[0].strip().lower() in {"total", "source", "standard"}:
        return None
    if len(values) < len(headers):
        values = values + [""] * (len(headers) - len(values))
    if len(values) > len(headers):
        values = values[: len(headers) - 1] + [" | ".join(values[len(headers) - 1 :])]
    row = dict(zip(headers, values))
    if not any(v.strip() for v in row.values()):
        return None
    return row


def candidate_key_field(headers: list[str], cohort: Cohort) -> str | None:
    if cohort.primary_key in headers:
        return cohort.primary_key
    low = {h.lower(): h for h in headers}
    for token in ("id", "code", "rule_id", "bom_id", "indicator_id", "form_code"):
        for key, original in low.items():
            if key == token or key.endswith("_" + token):
                return original
    return headers[0] if headers else None


def load_rows(cohort: Cohort) -> list[dict[str, str]]:
    base = PROJECT / cohort.slug / "research"
    if not base.exists():
        return []
    rows: list[dict[str, str]] = []
    seen = set()
    for path in sorted(base.glob("*data*.md")):
        for headers, raw_rows in markdown_tables(path):
            key_field = candidate_key_field(headers, cohort)
            if key_field is None or len(headers) < 3:
                continue
            for raw_line, values in raw_rows:
                item = normalize(headers, values, raw_line)
                if item is None:
                    continue
                key = item.get(cohort.primary_key) or item.get(key_field) or values[0]
                if not key or key.startswith("[") and "Table continues" in key:
                    continue
                fingerprint = (path.name, key, raw_line[:120])
                if fingerprint in seen:
                    continue
                seen.add(fingerprint)
                item["_source_file"] = path.relative_to(PROJECT).as_posix()
                item["_source_table_key"] = key_field
                item["_row_key"] = key
                rows.append(item)
    return rows


def all_fields(rows: list[dict[str, str]]) -> list[str]:
    fields: list[str] = []
    for row in rows:
        for key in row:
            if key.startswith("_"):
                continue
            if key not in fields:
                fields.append(key)
    return fields


def field_value(row: dict[str, str], field: str) -> str:
    if field in row:
        return row[field]
    lower = field.lower()
    for key, value in row.items():
        if key.lower() == lower:
            return value
    return ""


def counts(rows: list[dict[str, str]], field: str) -> Counter:
    c = Counter()
    for row in rows:
        value = field_value(row, field) or "[uncategorised]"
        c[value] += 1
    return c


def gap_count(rows: list[dict[str, str]]) -> int:
    return sum(1 for row in rows for value in row.values() if isinstance(value, str) and ("[GAP" in value or "[unverified" in value.lower()))


def source_tier_count(rows: list[dict[str, str]]) -> Counter:
    c = Counter()
    for row in rows:
        found = False
        for key, value in row.items():
            if "tier" in key.lower() and value:
                c[value] += 1
                found = True
        if not found:
            c["not explicit in row"] += 1
    return c


def import_status(row: dict[str, str], cohort: Cohort) -> str:
    text = " ".join(str(v) for v in row.values()).lower()
    if "[stub" in text:
        return "stub - do not import as active global setting"
    if "[gap" in text or "[unverified" in text:
        return "import with gap flag / curator review"
    if "placeholder" in text:
        return "import pointer only / needs narrative backfill"
    return "candidate global default"


def curation_priority(row: dict[str, str], cohort: Cohort) -> str:
    status = import_status(row, cohort)
    text = " ".join(str(v) for v in row.values()).lower()
    if "stub" in status:
        return "P0 fail-closed"
    if "placeholder" in status or "gap" in status:
        return "P1 curator review"
    if any(word in text for word in ("critical", "blocker", "mandatory", "required", "notifiable", "controlled")):
        return "P1 safety/compliance"
    return "P2 standard import"


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    for level, size in ((1, 17), (2, 13), (3, 11)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 78, 121)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade(cell, "1F4E79")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(8)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(headers)]):
            cells[idx].text = str(value)[:1600]
            for p in cells[idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)


def representative_fields(rows: list[dict[str, str]], cohort: Cohort) -> list[str]:
    fields = all_fields(rows)
    preferred = [cohort.primary_key, cohort.category_field]
    generic = ["name", "title", "description", "source_tier", "source_citations", "level_of_care_min", "cadre_min"]
    selected: list[str] = []
    for wanted in preferred + generic:
        for field in fields:
            if field not in selected and (field == wanted or field.lower() == wanted.lower() or wanted in field.lower()):
                selected.append(field)
                break
    for field in fields:
        if field not in selected:
            selected.append(field)
        if len(selected) >= 8:
            break
    return selected


def standard_rows(cohort: Cohort) -> list[list[str]]:
    rows = []
    for name in cohort.standards:
        maintainer, force, edition, use = STANDARDS.get(name, ("[source in cohort]", "See cohort source rows.", "See cohort.", "Used by this cohort."))
        rows.append([name, maintainer, force, edition, use])
    return rows


def write_reasoning(cohort: Cohort, rows: list[dict[str, str]]) -> Path:
    path = PROJECT / cohort.slug / "analysis" / "global-settings-critical-reasoning.md"
    top_groups = ", ".join(f"{k}: {v}" for k, v in counts(rows, cohort.category_field).most_common(8))
    text = f"""# Global Settings Critical Reasoning - {cohort.title}

**Date:** {DATE}

## Core question

Can this cohort be used as a Medic8 global-settings/defaults import source without hiding uncertainty?

## Evidence inventory

- Parsed rows: {len(rows)}
- Defaulting role: {cohort.preconfigure_scope}
- Facility confirmation: {cohort.facility_confirmation}
- Top groups: {top_groups}
- Gap/unverified cells retained: {gap_count(rows)}

## Argument map

| Conclusion | Evidence | Warrant | Countercase | Confidence |
|---|---|---|---|---|
| This cohort is suitable as an import planning artifact. | The local research files provide structured rows and source/gap markers. | Development and database teams need traceable defaults plus visible curator tasks. | Some rows are not production-ready and require review before activation. | Medium-high |
| Facility-specific values must remain separate from global defaults. | The onboarding contract distinguishes global seeds, facility scope decisions, extensions, mappings, and acceptance. | Values such as prices, stock counts, licences, staff, bank details, opening balances, and endpoint credentials cannot be safely guessed. | Over-automation could speed setup but create unsafe or false data. | High |
"""
    write(path, text)
    return path


def build_docx(cohort: Cohort, rows: list[dict[str, str]], reasoning: Path) -> Path:
    out = OUTPUT / f"{cohort.slug}-global-defaults-{VERSION}-{DATE}.docx"
    doc = Document()
    configure_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(f"{cohort.title} - Medic8 Global Defaults")
    run.bold = True
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph(f"Generated: {DATE}")
    doc.add_paragraph("Audience: Medic8 development and database teams preparing global settings, seed tables, import jobs, and curator worklists.")
    doc.add_paragraph("Evidence rule: every default must trace to a project row, local Medic8 file, or official standards source. Gaps remain explicit.")

    doc.add_heading("1. Purpose and Import Scope", level=1)
    doc.add_paragraph(cohort.purpose)
    add_table(doc, ["Field", "Definition"], [
        ["Import target", cohort.import_target],
        ["Preconfigure as default", cohort.preconfigure_scope],
        ["Facility confirmation required", cohort.facility_confirmation],
        ["Primary owner / signer", cohort.owner],
        ["Primary row key", cohort.primary_key],
        ["Grouping field", cohort.category_field],
    ])

    doc.add_heading("2. Standards and Defaults", level=1)
    doc.add_paragraph("The standards below define the canonical vocabulary or business rule family for this catalogue. The development team should preserve the canonical code, version/access date, source tier, and licensing caveats in global settings.")
    add_table(doc, ["Standard", "Maintainer", "Force / authority", "Edition or source", "Use in Medic8"], standard_rows(cohort))

    doc.add_heading("3. Dataset Summary", level=1)
    add_table(doc, ["Metric", "Value"], [
        ["Rows parsed", str(len(rows))],
        ["Gap/unverified cells retained", str(gap_count(rows))],
        ["Candidate global defaults", str(sum(1 for r in rows if import_status(r, cohort) == "candidate global default"))],
        ["Curator-review rows", str(sum(1 for r in rows if "review" in import_status(r, cohort)))],
        ["Stub/fail-closed rows", str(sum(1 for r in rows if "stub" in import_status(r, cohort)))],
    ])
    add_table(doc, ["Top group", "Rows"], [[k, str(v)] for k, v in counts(rows, cohort.category_field).most_common(12)])
    add_table(doc, ["Source tier", "Rows"], [[k, str(v)] for k, v in source_tier_count(rows).most_common(12)])

    doc.add_heading("4. Defaulting Rules for Development", level=1)
    for item in (
        "Import candidate global defaults inactive or draft unless the receiving table has a safe global-active semantic.",
        "Carry every source tier, source citation, code-system version, access date, and gap marker into staging.",
        "Never convert a facility-confirmation field into a global default.",
        "Fail closed for stub countries, unresolved licences, unresolved prices, missing BOMs, missing revenue accounts, and placeholder CDS narratives.",
        "Expose curator status in the import UI: candidate global default, review required, stub/fail-closed, or pointer-only.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. Facility Confirmation Boundary", level=1)
    doc.add_paragraph(cohort.facility_confirmation)
    doc.add_paragraph("The one-hour onboarding promise depends on collecting these non-guessable inputs before kickoff. During the onboarding call, the facility should confirm and sign off rather than discover missing data.")

    doc.add_heading("6. Representative Rows", level=1)
    fields = representative_fields(rows, cohort)
    display = [["curator_status", "curation_priority"] + fields]
    sample = []
    for row in rows[:35]:
        sample.append([import_status(row, cohort), curation_priority(row, cohort)] + [field_value(row, f) for f in fields])
    add_table(doc, display[0], sample)

    doc.add_heading("7. Gaps and Curator Worklist", level=1)
    gap_rows = [r for r in rows if import_status(r, cohort) != "candidate global default"][:25]
    if gap_rows:
        add_table(doc, ["row_key", "curator_status", "priority", "source_file"], [[r.get("_row_key", ""), import_status(r, cohort), curation_priority(r, cohort), r.get("_source_file", "")] for r in gap_rows])
    else:
        doc.add_paragraph("No gap/stub/placeholder rows were detected by the export heuristic. Curators must still review the source rows before production activation.")

    doc.add_heading("8. Acceptance Criteria", level=1)
    for item in (
        f"All rows from this cohort load into a staging table for `{cohort.import_target}`.",
        "The importer preserves raw source columns and adds curator status without overwriting source values.",
        "Global defaults are separated from facility-specific confirmations.",
        "Curator-review rows cannot become active production settings until signed off by the named owner.",
        "Foreign keys and cross-cohort references are checked after all global-settings workbooks are staged.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9. Critical Reasoning and Limits", level=1)
    doc.add_paragraph(f"Reasoning note: {reasoning.relative_to(PROJECT).as_posix()}")
    doc.add_paragraph("This document closes the research-to-import handoff. It does not certify production clinical safety, legal licensing, or local facility correctness.")

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def safe_sheet_name(name: str) -> str:
    cleaned = re.sub(r"[\[\]:*?/\\]", "-", name)
    return cleaned[:31] or "Sheet"


def write_sheet_rows(workbook, worksheet, headers: list[str], rows: list[list[str]], freeze: bool = True) -> None:
    header_fmt = workbook.add_format({"bold": True, "bg_color": "#1F4E79", "font_color": "white", "border": 1})
    normal = workbook.add_format({"text_wrap": True, "valign": "top", "border": 1})
    worksheet.write_row(0, 0, headers, header_fmt)
    for row_idx, row in enumerate(rows, 1):
        worksheet.write_row(row_idx, 0, [str(v)[:32000] for v in row], normal)
    if freeze:
        worksheet.freeze_panes(1, 0)
        worksheet.autofilter(0, 0, max(len(rows), 1), max(len(headers) - 1, 0))
    for idx, header in enumerate(headers):
        width = min(max(len(header) + 4, 14), 48)
        worksheet.set_column(idx, idx, width)


def build_xlsx(cohort: Cohort, rows: list[dict[str, str]]) -> Path:
    out = OUTPUT / f"{cohort.slug}-global-defaults-{VERSION}-{DATE}.xlsx"
    out.parent.mkdir(parents=True, exist_ok=True)
    workbook = xlsxwriter.Workbook(str(out))
    title_fmt = workbook.add_format({"bold": True, "font_size": 16, "font_color": "#1F4E79"})

    readme = workbook.add_worksheet("README")
    readme.write("A1", f"{cohort.title} - Medic8 Global Defaults", title_fmt)
    readme.write("A3", "Import target")
    readme.write("B3", cohort.import_target)
    readme.write("A4", "Preconfigure")
    readme.write("B4", cohort.preconfigure_scope)
    readme.write("A5", "Facility confirmation")
    readme.write("B5", cohort.facility_confirmation)
    readme.write("A6", "Owner")
    readme.write("B6", cohort.owner)
    readme.write("A7", "Rows")
    readme.write("B7", len(rows))
    readme.set_column("A:A", 26)
    readme.set_column("B:B", 120)

    write_sheet_rows(workbook, workbook.add_worksheet("Standards"), ["standard", "maintainer", "force", "edition_or_source", "use_in_medic8"], standard_rows(cohort))
    write_sheet_rows(workbook, workbook.add_worksheet("Summary"), ["metric", "value"], [
        ["rows", str(len(rows))],
        ["gap_unverified_cells", str(gap_count(rows))],
        ["candidate_global_defaults", str(sum(1 for r in rows if import_status(r, cohort) == "candidate global default"))],
        ["review_required_rows", str(sum(1 for r in rows if "review" in import_status(r, cohort)))],
        ["stub_fail_closed_rows", str(sum(1 for r in rows if "stub" in import_status(r, cohort)))],
    ], freeze=False)
    write_sheet_rows(workbook, workbook.add_worksheet("Groups"), ["group", "rows"], [[k, str(v)] for k, v in counts(rows, cohort.category_field).most_common()])

    fields = all_fields(rows)
    augmented = [
        "curator_status",
        "curation_priority",
        "preconfigure_scope",
        "facility_confirmation_required",
        "import_target_hint",
        "owner",
        "source_file",
        "row_key",
    ] + fields
    data_rows = []
    for row in rows:
        data_rows.append([
            import_status(row, cohort),
            curation_priority(row, cohort),
            cohort.preconfigure_scope,
            cohort.facility_confirmation,
            cohort.import_target,
            cohort.owner,
            row.get("_source_file", ""),
            row.get("_row_key", ""),
        ] + [field_value(row, f) for f in fields])
    write_sheet_rows(workbook, workbook.add_worksheet("Data"), augmented, data_rows)

    gap_rows = [r for r in rows if import_status(r, cohort) != "candidate global default"]
    write_sheet_rows(workbook, workbook.add_worksheet("Curator Worklist"), ["row_key", "status", "priority", "source_file", "notes"], [[r.get("_row_key", ""), import_status(r, cohort), curation_priority(r, cohort), r.get("_source_file", ""), "Review before production activation."] for r in gap_rows])

    workbook.close()
    return out


def build_master_docx(cohort_rows: dict[str, list[dict[str, str]]], artifacts: list[Path]) -> Path:
    out = OUTPUT / f"medic8-global-defaults-master-{VERSION}-{DATE}.docx"
    doc = Document()
    configure_doc(doc)
    title = doc.add_paragraph()
    run = title.add_run("Medic8 Global Defaults Handoff Pack")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 78, 121)
    doc.add_paragraph(f"Generated: {DATE}")
    doc.add_paragraph("Purpose: close the research project and hand a traceable defaults/import package to the Medic8 development and database teams.")

    doc.add_heading("1. Handoff Boundary", level=1)
    doc.add_paragraph("The research corpus defines global defaults, standards, gap flags, and curator worklists. Development/database teams should stage these into global-setting tables, then curate, sort, validate, and activate them through product-owned import jobs.")
    doc.add_paragraph("The one-hour onboarding promise remains conditional: global defaults can be preconfigured, but facility-specific inputs such as staff, licences, prices, bank details, opening balances, stock counts, insurer contracts, endpoint credentials, and patient migrations must be supplied or confirmed by the facility.")

    doc.add_heading("2. Cohort Inventory", level=1)
    rows = []
    for cohort in COHORTS:
        data = cohort_rows[cohort.slug]
        rows.append([cohort.slug, cohort.title, str(len(data)), cohort.import_target, cohort.owner, cohort.preconfigure_scope, cohort.facility_confirmation])
    add_table(doc, ["cohort", "title", "rows", "import target", "owner", "preconfigure", "facility confirmation"], rows)

    doc.add_heading("3. Productization Gates", level=1)
    for item in (
        "Create staging tables/import jobs that preserve raw source values and add curator status.",
        "Map each cohort to global settings before facility-specific tenant rows are created.",
        "Fail closed for stubs, gaps, unresolved licences, unresolved BOMs, missing revenue accounts, and placeholder CDS narratives.",
        "Run row-level curator sign-off by the owner named in each workbook.",
        "Run timed acceptance tests for all six canonical tenant blueprints before using the one-hour promise without qualification.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Artifact Index", level=1)
    add_table(doc, ["artifact"], [[p.relative_to(PROJECT).as_posix()] for p in artifacts])
    doc.save(out)
    return out


def build_master_xlsx(cohort_rows: dict[str, list[dict[str, str]]]) -> Path:
    out = OUTPUT / f"medic8-global-defaults-import-workbook-{VERSION}-{DATE}.xlsx"
    workbook = xlsxwriter.Workbook(str(out))
    inventory_rows = []
    for cohort in COHORTS:
        rows = cohort_rows[cohort.slug]
        inventory_rows.append([
            cohort.slug,
            cohort.title,
            len(rows),
            cohort.import_target,
            cohort.preconfigure_scope,
            cohort.facility_confirmation,
            cohort.owner,
            gap_count(rows),
        ])
    write_sheet_rows(workbook, workbook.add_worksheet("Cohort Inventory"), ["cohort", "title", "rows", "import_target", "preconfigure_scope", "facility_confirmation", "owner", "gap_unverified_cells"], inventory_rows)

    standards_seen = []
    for cohort in COHORTS:
        for standard in cohort.standards:
            if standard not in standards_seen:
                standards_seen.append(standard)
    write_sheet_rows(workbook, workbook.add_worksheet("Standards Registry"), ["standard", "maintainer", "force", "edition_or_source", "use"], [[s] + list(STANDARDS.get(s, ("[source in cohort]", "See cohort rows", "See cohort", "Used by cohort"))) for s in standards_seen])

    worklist = []
    for cohort in COHORTS:
        for row in cohort_rows[cohort.slug]:
            status = import_status(row, cohort)
            if status != "candidate global default":
                worklist.append([cohort.slug, row.get("_row_key", ""), status, curation_priority(row, cohort), row.get("_source_file", ""), cohort.owner])
    write_sheet_rows(workbook, workbook.add_worksheet("Global Curator Worklist"), ["cohort", "row_key", "status", "priority", "source_file", "owner"], worklist)

    for cohort in COHORTS:
        rows = cohort_rows[cohort.slug]
        fields = all_fields(rows)
        sheet_headers = ["curator_status", "priority", "source_file", "row_key"] + fields[:35]
        sheet_rows = []
        for row in rows:
            sheet_rows.append([import_status(row, cohort), curation_priority(row, cohort), row.get("_source_file", ""), row.get("_row_key", "")] + [field_value(row, f) for f in fields[:35]])
        write_sheet_rows(workbook, workbook.add_worksheet(safe_sheet_name(cohort.slug)), sheet_headers, sheet_rows)

    workbook.close()
    return out


def write_manifest(artifacts: list[Path], cohort_rows: dict[str, list[dict[str, str]]]) -> Path:
    lines = [
        "# Medic8 Global Settings Deliverables Manifest",
        "",
        f"**Generated:** {DATE}",
        f"**Version:** {VERSION}",
        "",
        "## Purpose",
        "",
        "Closeout handoff package for Medic8 development and database teams. These DOCX and XLSX artifacts define source-traceable global defaults, standards, import targets, facility-confirmation boundaries, and curator worklists.",
        "",
        "## Row Counts",
        "",
    ]
    for cohort in COHORTS:
        lines.append(f"- {cohort.slug}: {len(cohort_rows[cohort.slug])} rows parsed")
    lines.extend(["", "## Artifacts", ""])
    for artifact in artifacts:
        lines.append(f"- `{artifact.relative_to(PROJECT).as_posix()}`")
    lines.extend([
        "",
        "## Evidence Note",
        "",
        "Gap, stub, unverified, and placeholder values are retained and converted into curator worklist items. They are not production-ready global settings until reviewed.",
    ])
    path = OUTPUT / "manifest.md"
    write(path, "\n".join(lines) + "\n")
    return path


def validate_office(path: Path) -> bool:
    try:
        with zipfile.ZipFile(path) as zf:
            return zf.testzip() is None
    except zipfile.BadZipFile:
        return False


def update_status() -> None:
    path = PROJECT / "PROJECT-STATUS.md"
    text = read(path)
    addition = f"""

## Phase 9 - Closeout deliverables for development/database handoff (added {DATE})

**Output family:** `05-output/medic8-global-settings/`

**Purpose:** regenerate Word and Excel artifacts as a v2 closeout package that clearly defines global defaults, standards, import targets, facility-confirmation boundaries, curator worklists, and source/gap discipline for Medic8 development and database teams.

**Artifacts:** per-cohort DOCX/XLSX files for all active cohorts, plus `medic8-global-defaults-master-{VERSION}-{DATE}.docx`, `medic8-global-defaults-import-workbook-{VERSION}-{DATE}.xlsx`, and `manifest.md`.

**Handoff rule:** these files define staging/default candidates. Development must preserve source/gap fields, fail closed for stubs/placeholders, and require curator sign-off before production activation.

**Validation:** generated Office files are ZIP-integrity checked by `scripts/generate_medic8_global_settings_outputs.py`; run `python -m engine validate healthcare-app-clinical-data` before release.
"""
    if "Phase 9 - Closeout deliverables" not in text:
        write(path, text.rstrip() + addition)


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    EXPORT.mkdir(parents=True, exist_ok=True)
    cohort_rows: dict[str, list[dict[str, str]]] = {}
    artifacts: list[Path] = []
    for cohort in COHORTS:
        rows = load_rows(cohort)
        cohort_rows[cohort.slug] = rows
        reasoning = write_reasoning(cohort, rows)
        docx = build_docx(cohort, rows, reasoning)
        xlsx = build_xlsx(cohort, rows)
        artifacts.extend([reasoning, docx, xlsx])

    master_docx = build_master_docx(cohort_rows, artifacts)
    master_xlsx = build_master_xlsx(cohort_rows)
    artifacts.extend([master_docx, master_xlsx])
    manifest = write_manifest(artifacts, cohort_rows)
    artifacts.append(manifest)
    update_status()

    for artifact in artifacts:
        if artifact.suffix in {".docx", ".xlsx"} or artifact.name == "manifest.md":
            shutil.copy2(artifact, EXPORT / artifact.name)

    invalid = [p for p in artifacts if p.suffix in {".docx", ".xlsx"} and not validate_office(p)]
    if invalid:
        raise SystemExit("Invalid Office artifacts: " + ", ".join(str(p) for p in invalid))

    print("Generated Medic8 global settings closeout package")
    for cohort in COHORTS:
        print(f"{cohort.slug}: {len(cohort_rows[cohort.slug])} rows")
    package_files = [p for p in artifacts if p.suffix in {".docx", ".xlsx"} or p.name == "manifest.md"]
    print(f"Output package files: {len(package_files)}")
    print(f"Analysis files: {len(artifacts) - len(package_files)}")


if __name__ == "__main__":
    main()
